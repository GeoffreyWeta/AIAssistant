import os
import fitz  # PyMuPDF
import logging
from typing import List, Tuple

from langchain.docstore.document import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma

from app.vectorstore.embedding import get_embedding_model
from app.Config.config import settings

logger = logging.getLogger(__name__)

SUPPORTED_EXTS = {".pdf", ".docx", ".txt"}

def _extract_texts_with_pages_pdf(file_path: str) -> List[Tuple[int, str]]:
    texts: List[Tuple[int, str]] = []
    with fitz.open(file_path) as doc:
        for i, page in enumerate(doc, start=1):
            texts.append((i, page.get_text()))
    return texts

def _extract_texts_docx(file_path: str) -> List[Tuple[int, str]]:
    from docx import Document as Docx
    d = Docx(file_path)
    text = "\n".join(p.text for p in d.paragraphs)
    # treat as one page for now
    return [(1, text)]

def _extract_texts_txt(file_path: str) -> List[Tuple[int, str]]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [(1, text)]

def extract_pages(file_path: str) -> List[Tuple[int, str]]:
    ext = os.path.splitext(file_path)[1].lower()
    if ext not in SUPPORTED_EXTS:
        raise ValueError(f"Unsupported file type: {ext}")
    if ext == ".pdf":
        return _extract_texts_with_pages_pdf(file_path)
    if ext == ".docx":
        return _extract_texts_docx(file_path)
    return _extract_texts_txt(file_path)

def load_and_vectorize(
    file_path: str,
    doc_id: str,
    filename: str,
    persist_directory: str,
    summarize: bool = False,
):
    """
    Extracts text per page, chunks while preserving page metadata,
    embeds and persists to a Chroma store located at persist_directory.
    """
    pages = extract_pages(file_path)
    total_chars = sum(len(t) for _, t in pages)
    if total_chars < 10:
        raise ValueError("No text found in document")

    # Split per page to preserve page numbers
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=120)
    docs: List[Document] = []
    for page_no, text in pages:
        if not text.strip():
            continue
        base_doc = Document(page_content=text, metadata={"page": page_no})
        for chunk in splitter.split_documents([base_doc]):
            # carry page forward and enrich metadata
            meta = {
                "doc_id": doc_id,
                "filename": filename,
                "page": chunk.metadata.get("page", page_no),
            }
            docs.append(Document(page_content=chunk.page_content, metadata=meta))

    if not docs:
        raise ValueError("No chunks produced from document")

    # Optional summarization pipeline can be added later, default is off for speed

    os.makedirs(persist_directory, exist_ok=True)
    embedding = get_embedding_model()
    db = Chroma.from_documents(documents=docs, embedding=embedding, persist_directory=persist_directory)
    db.persist()

    logger.info(f"Vectorized {filename}, chunks: {len(docs)}, stored at {persist_directory}")
    return {"chunks": len(docs)}
